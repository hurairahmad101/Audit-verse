import io
import os
import logging
from typing import List, Optional
from datetime import datetime, timedelta
from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func
from pydantic import BaseModel

from ....models import (
    AuditReport, AuditBoardPack, AuditEngagement, AuditFinding,
    AuditRecommendation, AuditActionPlan, AuditManagementResponse,
    AuditPlan, AuditPlanItem,
    PBCListItem, GRCUser, AuditableEntity, Risk, get_db,
    RegulatoryChange, RegulatoryImpactAssessment, RegulatoryFeedItem
)
from ....routers.auth_router import require_auth, get_user_tenants, get_user_primary_tenant

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/reporting", tags=["Audit - Reporting"])


class ReportCreate(BaseModel):
    engagement_id: int
    title: str
    report_type: Optional[str] = "engagement_report"
    executive_summary: Optional[str] = None
    opinion: Optional[str] = None
    opinion_narrative: Optional[str] = None
    scope_summary: Optional[str] = None


class ReportUpdate(BaseModel):
    title: Optional[str] = None
    executive_summary: Optional[str] = None
    opinion: Optional[str] = None
    opinion_narrative: Optional[str] = None
    scope_summary: Optional[str] = None
    status: Optional[str] = None


class BoardPackCreate(BaseModel):
    title: str
    period: Optional[str] = None
    engagement_ids: Optional[list] = []
    executive_summary: Optional[str] = None


class PBCItemCreate(BaseModel):
    engagement_id: int
    document_name: str
    description: Optional[str] = None
    category: Optional[str] = None
    requested_by: Optional[str] = None
    assigned_to_id: Optional[int] = None
    due_date: Optional[datetime] = None


class PBCItemUpdate(BaseModel):
    status: Optional[str] = None
    assigned_to_id: Optional[int] = None
    due_date: Optional[datetime] = None
    evidence_id: Optional[int] = None
    notes: Optional[str] = None


def serialize_report(r: AuditReport) -> dict:
    return {
        "id": r.id,
        "tenant_id": r.tenant_id,
        "engagement_id": r.engagement_id,
        "title": r.title,
        "report_type": r.report_type,
        "executive_summary": r.executive_summary,
        "opinion": r.opinion,
        "opinion_narrative": r.opinion_narrative,
        "scope_summary": r.scope_summary,
        "findings_summary": r.findings_summary,
        "recommendations_summary": r.recommendations_summary,
        "status": r.status,
        "ai_generated": r.ai_generated,
        "issued_date": r.issued_date.isoformat() if r.issued_date else None,
        "issued_by_id": r.issued_by_id,
        "created_at": r.created_at.isoformat() if r.created_at else None,
    }


@router.get("/reports")
def list_reports(
    engagement_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"reports": [], "total": 0}
    
    query = db.query(AuditReport).filter(AuditReport.tenant_id.in_(user_tenants))
    if engagement_id:
        query = query.filter(AuditReport.engagement_id == engagement_id)
    
    reports = query.order_by(AuditReport.created_at.desc()).all()
    return {"reports": [serialize_report(r) for r in reports], "total": len(reports)}


@router.post("/reports", status_code=status.HTTP_201_CREATED)
def create_report(
    data: ReportCreate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    eng = db.query(AuditEngagement).filter(
        AuditEngagement.id == data.engagement_id,
        AuditEngagement.tenant_id.in_(user_tenants)
    ).first()
    if not eng:
        raise HTTPException(status_code=404, detail="Engagement not found")
    
    findings = db.query(AuditFinding).filter(
        AuditFinding.engagement_id == data.engagement_id
    ).all()
    
    findings_summary = {
        "total": len(findings),
        "by_severity": {},
        "by_status": {},
    }
    for f in findings:
        findings_summary["by_severity"][f.severity] = findings_summary["by_severity"].get(f.severity, 0) + 1
        findings_summary["by_status"][f.status] = findings_summary["by_status"].get(f.status, 0) + 1
    
    recs = db.query(AuditRecommendation).join(AuditFinding).filter(
        AuditFinding.engagement_id == data.engagement_id
    ).all()
    
    recs_summary = {
        "total": len(recs),
        "by_priority": {},
        "by_status": {},
    }
    for r in recs:
        recs_summary["by_priority"][r.priority] = recs_summary["by_priority"].get(r.priority, 0) + 1
        recs_summary["by_status"][r.status] = recs_summary["by_status"].get(r.status, 0) + 1
    
    report = AuditReport(
        tenant_id=eng.tenant_id,
        engagement_id=data.engagement_id,
        title=data.title,
        report_type=data.report_type,
        executive_summary=data.executive_summary,
        opinion=data.opinion,
        opinion_narrative=data.opinion_narrative,
        scope_summary=data.scope_summary,
        findings_summary=findings_summary,
        recommendations_summary=recs_summary,
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    return serialize_report(report)


@router.put("/reports/{report_id}")
def update_report(
    report_id: int,
    data: ReportUpdate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    report = db.query(AuditReport).filter(
        AuditReport.id == report_id,
        AuditReport.tenant_id.in_(user_tenants)
    ).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    for field, value in data.dict(exclude_unset=True).items():
        setattr(report, field, value)
    
    if data.status == "issued":
        report.issued_date = datetime.utcnow()
        report.issued_by_id = current_user.id
    
    report.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(report)
    return serialize_report(report)


def _get_full_report_data(report_id: int, db: Session, user_tenants: list) -> dict:
    report = db.query(AuditReport).filter(
        AuditReport.id == report_id,
        AuditReport.tenant_id.in_(user_tenants)
    ).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    engagement = db.query(AuditEngagement).filter(
        AuditEngagement.id == report.engagement_id
    ).first()

    entity = None
    if engagement and engagement.auditable_entity_id:
        entity = db.query(AuditableEntity).filter(
            AuditableEntity.id == engagement.auditable_entity_id
        ).first()

    lead_auditor = None
    if engagement and engagement.lead_auditor_id:
        lead_auditor = db.query(GRCUser).filter(GRCUser.id == engagement.lead_auditor_id).first()

    findings = db.query(AuditFinding).filter(
        AuditFinding.engagement_id == report.engagement_id
    ).order_by(AuditFinding.created_at).all()

    findings_data = []
    for f in findings:
        recs = db.query(AuditRecommendation).filter(
            AuditRecommendation.finding_id == f.id
        ).all()

        mgmt_responses = db.query(AuditManagementResponse).filter(
            AuditManagementResponse.finding_id == f.id
        ).all()

        recs_data = []
        for rec in recs:
            action_plans = db.query(AuditActionPlan).filter(
                AuditActionPlan.recommendation_id == rec.id
            ).all()
            rec_owner = db.query(GRCUser).filter(GRCUser.id == rec.owner_id).first() if rec.owner_id else None
            recs_data.append({
                "id": rec.id,
                "title": rec.title,
                "description": rec.description,
                "priority": rec.priority,
                "status": rec.status,
                "owner": rec_owner.display_name or rec_owner.username if rec_owner else None,
                "due_date": rec.due_date.isoformat() if rec.due_date else None,
                "action_plans": [{
                    "id": ap.id,
                    "milestone": ap.milestone,
                    "description": ap.description,
                    "status": ap.status,
                    "due_date": ap.due_date.isoformat() if ap.due_date else None,
                    "completed_date": ap.completed_date.isoformat() if ap.completed_date else None,
                } for ap in action_plans],
            })

        finding_owner = db.query(GRCUser).filter(GRCUser.id == f.owner_id).first() if f.owner_id else None
        findings_data.append({
            "id": f.id,
            "finding_number": f.finding_number,
            "title": f.title,
            "condition": f.condition,
            "criteria": f.criteria,
            "cause": f.cause,
            "effect": f.effect,
            "root_cause_category": f.root_cause_category,
            "severity": f.severity,
            "status": f.status,
            "theme": f.theme,
            "owner": finding_owner.display_name or finding_owner.username if finding_owner else None,
            "due_date": f.due_date.isoformat() if f.due_date else None,
            "recommendations": recs_data,
            "management_responses": [{
                "id": mr.id,
                "response_type": mr.response_type,
                "response_text": mr.response_text,
                "action_plan": mr.action_plan,
                "target_date": mr.target_date.isoformat() if mr.target_date else None,
                "responded_at": mr.responded_at.isoformat() if mr.responded_at else None,
            } for mr in mgmt_responses],
        })

    issued_by = db.query(GRCUser).filter(GRCUser.id == report.issued_by_id).first() if report.issued_by_id else None

    return {
        "id": report.id,
        "title": report.title,
        "report_type": report.report_type,
        "status": report.status,
        "opinion": report.opinion,
        "opinion_narrative": report.opinion_narrative,
        "executive_summary": report.executive_summary,
        "scope_summary": report.scope_summary,
        "findings_summary": report.findings_summary,
        "recommendations_summary": report.recommendations_summary,
        "ai_generated": report.ai_generated,
        "issued_date": report.issued_date.isoformat() if report.issued_date else None,
        "issued_by": issued_by.display_name or issued_by.username if issued_by else None,
        "created_at": report.created_at.isoformat() if report.created_at else None,
        "updated_at": report.updated_at.isoformat() if report.updated_at else None,
        "engagement": {
            "id": engagement.id,
            "title": engagement.title,
            "description": engagement.description,
            "engagement_type": engagement.engagement_type,
            "status": engagement.status,
            "scope": engagement.scope,
            "objectives": engagement.objectives,
            "methodology": engagement.methodology,
            "planned_start": engagement.planned_start.isoformat() if engagement.planned_start else None,
            "planned_end": engagement.planned_end.isoformat() if engagement.planned_end else None,
            "actual_start": engagement.actual_start.isoformat() if engagement.actual_start else None,
            "actual_end": engagement.actual_end.isoformat() if engagement.actual_end else None,
            "lead_auditor": lead_auditor.display_name or lead_auditor.username if lead_auditor else None,
            "entity_name": entity.name if entity else None,
        } if engagement else None,
        "findings": findings_data,
    }


@router.get("/reports/{report_id}/full")
def get_full_report(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        raise HTTPException(status_code=403, detail="No tenant access")
    return _get_full_report_data(report_id, db, user_tenants)


@router.get("/reports/{report_id}/export/pdf")
def export_report_pdf(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        raise HTTPException(status_code=403, detail="No tenant access")

    data = _get_full_report_data(report_id, db, user_tenants)

    try:
        from fpdf import FPDF

        class AuditPDF(FPDF):
            def header(self):
                self.set_font('Helvetica', 'B', 10)
                self.set_text_color(100, 100, 100)
                self.cell(0, 8, data["title"], align='R')
                self.ln(10)
                self.set_draw_color(59, 130, 246)
                self.line(10, self.get_y(), 200, self.get_y())
                self.ln(5)

            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')

            def section_title(self, title):
                self.set_font('Helvetica', 'B', 14)
                self.set_text_color(30, 58, 138)
                self.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
                self.set_draw_color(59, 130, 246)
                self.line(10, self.get_y(), 80, self.get_y())
                self.ln(4)

            def sub_title(self, title):
                self.set_font('Helvetica', 'B', 11)
                self.set_text_color(50, 50, 50)
                self.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
                self.ln(1)

            def body_text(self, text):
                if not text:
                    return
                self.set_font('Helvetica', '', 10)
                self.set_text_color(60, 60, 60)
                self.multi_cell(0, 5, text)
                self.ln(3)

            def badge(self, text, r, g, b):
                self.set_font('Helvetica', 'B', 9)
                self.set_fill_color(r, g, b)
                self.set_text_color(255, 255, 255)
                w = self.get_string_width(text) + 8
                self.cell(w, 7, text, fill=True, new_x="END")
                self.cell(3)

        pdf = AuditPDF()
        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=20)

        pdf.add_page()
        pdf.ln(30)
        pdf.set_font('Helvetica', 'B', 28)
        pdf.set_text_color(30, 41, 59)
        pdf.multi_cell(0, 12, data["title"], align='C')
        pdf.ln(10)

        pdf.set_font('Helvetica', '', 14)
        pdf.set_text_color(100, 116, 139)
        report_type_label = (data.get("report_type") or "").replace("_", " ").title()
        pdf.cell(0, 8, report_type_label, align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

        if data.get("engagement"):
            eng = data["engagement"]
            if eng.get("entity_name"):
                pdf.cell(0, 8, eng["entity_name"], align='C', new_x="LMARGIN", new_y="NEXT")

        pdf.ln(15)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(60, pdf.get_y(), 150, pdf.get_y())
        pdf.ln(10)

        pdf.set_font('Helvetica', '', 11)
        pdf.set_text_color(80, 80, 80)
        status_label = (data.get("status") or "draft").upper()
        opinion_label = (data.get("opinion") or "").replace("_", " ").title()
        pdf.cell(0, 7, f"Status: {status_label}", align='C', new_x="LMARGIN", new_y="NEXT")
        if opinion_label:
            pdf.cell(0, 7, f"Opinion: {opinion_label}", align='C', new_x="LMARGIN", new_y="NEXT")
        if data.get("issued_date"):
            pdf.cell(0, 7, f"Issued: {data['issued_date'][:10]}", align='C', new_x="LMARGIN", new_y="NEXT")
        if data.get("issued_by"):
            pdf.cell(0, 7, f"Issued by: {data['issued_by']}", align='C', new_x="LMARGIN", new_y="NEXT")

        eng = data.get("engagement") or {}
        if eng.get("lead_auditor"):
            pdf.cell(0, 7, f"Lead Auditor: {eng['lead_auditor']}", align='C', new_x="LMARGIN", new_y="NEXT")

        pdf.ln(10)
        pdf.set_font('Helvetica', 'I', 9)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(0, 7, f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}", align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 7, "AuditVerse.AI", align='C', new_x="LMARGIN", new_y="NEXT")

        pdf.add_page()
        pdf.section_title("TABLE OF CONTENTS")
        toc_items = ["1. Executive Summary", "2. Engagement Overview", "3. Scope & Methodology"]
        if data.get("opinion") or data.get("opinion_narrative"):
            toc_items.append("4. Auditor's Opinion")
        toc_items.append(f"{'5' if len(toc_items) == 4 else '4'}. Detailed Findings")
        toc_items.append(f"{'6' if len(toc_items) == 5 else '5'}. Recommendations & Action Plans")
        for item in toc_items:
            pdf.set_font('Helvetica', '', 11)
            pdf.set_text_color(60, 60, 60)
            pdf.cell(0, 8, item, new_x="LMARGIN", new_y="NEXT")

        pdf.add_page()
        pdf.section_title("1. EXECUTIVE SUMMARY")
        pdf.body_text(data.get("executive_summary") or "No executive summary provided.")

        pdf.section_title("2. ENGAGEMENT OVERVIEW")
        if eng:
            info_lines = [
                ("Engagement", eng.get("title", "")),
                ("Type", (eng.get("engagement_type") or "").replace("_", " ").title()),
                ("Entity", eng.get("entity_name") or "N/A"),
                ("Lead Auditor", eng.get("lead_auditor") or "N/A"),
                ("Period", f"{(eng.get('planned_start') or '')[:10]} to {(eng.get('planned_end') or '')[:10]}" if eng.get('planned_start') else "N/A"),
                ("Status", (eng.get("status") or "").replace("_", " ").title()),
            ]
            for label, value in info_lines:
                pdf.set_font('Helvetica', 'B', 10)
                pdf.set_text_color(60, 60, 60)
                pdf.cell(45, 7, f"{label}:")
                pdf.set_font('Helvetica', '', 10)
                pdf.cell(0, 7, value, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
            if eng.get("description"):
                pdf.sub_title("Description")
                pdf.body_text(eng["description"])
            if eng.get("objectives"):
                pdf.sub_title("Objectives")
                pdf.body_text(eng["objectives"])

        pdf.section_title("3. SCOPE & METHODOLOGY")
        pdf.body_text(data.get("scope_summary") or eng.get("scope") or "No scope information provided.")
        if eng.get("methodology"):
            pdf.sub_title("Methodology")
            pdf.body_text(eng["methodology"])

        if data.get("opinion") or data.get("opinion_narrative"):
            pdf.add_page()
            pdf.section_title("4. AUDITOR'S OPINION")
            if data.get("opinion"):
                sev_colors = {
                    "satisfactory": (34, 197, 94),
                    "needs_improvement": (245, 158, 11),
                    "unsatisfactory": (239, 68, 68),
                    "advisory": (59, 130, 246),
                }
                color = sev_colors.get(data["opinion"], (100, 100, 100))
                pdf.badge(data["opinion"].replace("_", " ").upper(), *color)
                pdf.ln(8)
            pdf.body_text(data.get("opinion_narrative") or "")

        findings = data.get("findings", [])
        pdf.add_page()
        section_num = 5 if (data.get("opinion") or data.get("opinion_narrative")) else 4
        pdf.section_title(f"{section_num}. DETAILED FINDINGS")

        if not findings:
            pdf.body_text("No findings were identified during this engagement.")
        else:
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(80, 80, 80)
            pdf.cell(0, 7, f"Total findings: {len(findings)}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)

            for i, finding in enumerate(findings):
                if pdf.get_y() > 230:
                    pdf.add_page()
                pdf.set_font('Helvetica', 'B', 12)
                pdf.set_text_color(30, 41, 59)
                finding_label = finding.get("finding_number") or f"F-{i+1}"
                pdf.cell(0, 8, f"{finding_label}: {finding['title']}", new_x="LMARGIN", new_y="NEXT")

                sev = finding.get("severity", "medium")
                sev_colors_f = {"critical": (220, 38, 38), "high": (234, 88, 12), "medium": (245, 158, 11), "low": (34, 197, 94)}
                color = sev_colors_f.get(sev, (100, 100, 100))
                pdf.badge(sev.upper(), *color)
                status_f = finding.get("status", "open")
                pdf.badge(status_f.upper(), 100, 116, 139)
                pdf.ln(8)

                for field_label, field_key in [("Condition", "condition"), ("Criteria", "criteria"), ("Cause", "cause"), ("Effect", "effect")]:
                    if finding.get(field_key):
                        pdf.sub_title(field_label)
                        pdf.body_text(finding[field_key])

                if finding.get("owner"):
                    pdf.set_font('Helvetica', 'I', 9)
                    pdf.set_text_color(120, 120, 120)
                    pdf.cell(0, 6, f"Owner: {finding['owner']}", new_x="LMARGIN", new_y="NEXT")

                for mr in finding.get("management_responses", []):
                    pdf.sub_title("Management Response")
                    resp_type = (mr.get("response_type") or "").replace("_", " ").title()
                    pdf.set_font('Helvetica', 'I', 9)
                    pdf.set_text_color(100, 100, 100)
                    pdf.cell(0, 6, f"Response: {resp_type}", new_x="LMARGIN", new_y="NEXT")
                    pdf.body_text(mr.get("response_text") or "")
                    if mr.get("action_plan"):
                        pdf.body_text(f"Action Plan: {mr['action_plan']}")

                pdf.ln(5)
                pdf.set_draw_color(220, 220, 220)
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                pdf.ln(5)

        pdf.add_page()
        pdf.section_title(f"{section_num + 1}. RECOMMENDATIONS & ACTION PLANS")
        all_recs = []
        for f in findings:
            for rec in f.get("recommendations", []):
                rec["finding_title"] = f["title"]
                all_recs.append(rec)

        if not all_recs:
            pdf.body_text("No recommendations recorded.")
        else:
            for rec in all_recs:
                if pdf.get_y() > 240:
                    pdf.add_page()
                pdf.set_font('Helvetica', 'B', 11)
                pdf.set_text_color(30, 41, 59)
                pdf.cell(0, 8, rec["title"], new_x="LMARGIN", new_y="NEXT")

                pri = rec.get("priority", "medium")
                pri_colors = {"critical": (220, 38, 38), "high": (234, 88, 12), "medium": (245, 158, 11), "low": (34, 197, 94)}
                color = pri_colors.get(pri, (100, 100, 100))
                pdf.badge(pri.upper(), *color)
                pdf.badge(rec.get("status", "open").upper(), 100, 116, 139)
                pdf.ln(6)

                pdf.set_font('Helvetica', 'I', 9)
                pdf.set_text_color(120, 120, 120)
                pdf.cell(0, 6, f"Related Finding: {rec['finding_title']}", new_x="LMARGIN", new_y="NEXT")

                if rec.get("description"):
                    pdf.body_text(rec["description"])
                if rec.get("owner"):
                    pdf.body_text(f"Owner: {rec['owner']}")
                if rec.get("due_date"):
                    pdf.body_text(f"Due: {rec['due_date'][:10]}")

                for ap in rec.get("action_plans", []):
                    pdf.set_font('Helvetica', '', 9)
                    pdf.set_text_color(80, 80, 80)
                    status_sym = "✓" if ap.get("status") == "completed" else "○"
                    pdf.cell(0, 6, f"  {status_sym} {ap['milestone']} ({ap.get('status', 'pending')})", new_x="LMARGIN", new_y="NEXT")

                pdf.ln(4)

        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        filename = data["title"].replace(" ", "_").replace("/", "_") + ".pdf"
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        logger.error(f"PDF export error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


@router.get("/reports/{report_id}/export/docx")
def export_report_docx(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        raise HTTPException(status_code=403, detail="No tenant access")

    data = _get_full_report_data(report_id, db, user_tenants)

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(60, 60, 60)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(data["title"])
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(30, 41, 59)

        report_type_label = (data.get("report_type") or "").replace("_", " ").title()
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(report_type_label)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(100, 116, 139)

        eng = data.get("engagement") or {}
        if eng.get("entity_name"):
            entity_para = doc.add_paragraph()
            entity_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            entity_run = entity_para.add_run(eng["entity_name"])
            entity_run.font.size = Pt(12)
            entity_run.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph()
        meta_items = []
        if data.get("status"):
            meta_items.append(f"Status: {data['status'].upper()}")
        if data.get("opinion"):
            meta_items.append(f"Opinion: {data['opinion'].replace('_', ' ').title()}")
        if data.get("issued_date"):
            meta_items.append(f"Issued: {data['issued_date'][:10]}")
        if data.get("issued_by"):
            meta_items.append(f"Issued by: {data['issued_by']}")
        if eng.get("lead_auditor"):
            meta_items.append(f"Lead Auditor: {eng['lead_auditor']}")

        for item in meta_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(item)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_page_break()

        h = doc.add_heading('1. Executive Summary', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(30, 58, 138)
        doc.add_paragraph(data.get("executive_summary") or "No executive summary provided.")

        h = doc.add_heading('2. Engagement Overview', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(30, 58, 138)

        if eng:
            info_lines = [
                ("Engagement", eng.get("title", "")),
                ("Type", (eng.get("engagement_type") or "").replace("_", " ").title()),
                ("Entity", eng.get("entity_name") or "N/A"),
                ("Lead Auditor", eng.get("lead_auditor") or "N/A"),
                ("Period", f"{(eng.get('planned_start') or '')[:10]} to {(eng.get('planned_end') or '')[:10]}" if eng.get('planned_start') else "N/A"),
            ]
            for label, value in info_lines:
                p = doc.add_paragraph()
                r = p.add_run(f"{label}: ")
                r.bold = True
                p.add_run(value)

            if eng.get("description"):
                doc.add_heading("Description", level=2)
                doc.add_paragraph(eng["description"])
            if eng.get("objectives"):
                doc.add_heading("Objectives", level=2)
                doc.add_paragraph(eng["objectives"])

        h = doc.add_heading('3. Scope & Methodology', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(30, 58, 138)
        doc.add_paragraph(data.get("scope_summary") or eng.get("scope") or "No scope information provided.")
        if eng.get("methodology"):
            doc.add_heading("Methodology", level=2)
            doc.add_paragraph(eng["methodology"])

        section_num = 4
        if data.get("opinion") or data.get("opinion_narrative"):
            h = doc.add_heading(f'{section_num}. Auditor\'s Opinion', level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(30, 58, 138)
            if data.get("opinion"):
                p = doc.add_paragraph()
                r = p.add_run(f"Opinion: {data['opinion'].replace('_', ' ').upper()}")
                r.bold = True
                sev_colors_o = {
                    "satisfactory": RGBColor(34, 197, 94),
                    "needs_improvement": RGBColor(245, 158, 11),
                    "unsatisfactory": RGBColor(239, 68, 68),
                }
                r.font.color.rgb = sev_colors_o.get(data["opinion"], RGBColor(100, 100, 100))
            if data.get("opinion_narrative"):
                doc.add_paragraph(data["opinion_narrative"])
            section_num += 1

        findings = data.get("findings", [])
        h = doc.add_heading(f'{section_num}. Detailed Findings', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(30, 58, 138)

        if not findings:
            doc.add_paragraph("No findings were identified during this engagement.")
        else:
            doc.add_paragraph(f"Total findings: {len(findings)}")
            for i, finding in enumerate(findings):
                finding_label = finding.get("finding_number") or f"F-{i+1}"
                fh = doc.add_heading(f"{finding_label}: {finding['title']}", level=2)
                p = doc.add_paragraph()
                r = p.add_run(f"Severity: {finding.get('severity', 'medium').upper()}")
                r.bold = True
                sev_c = {"critical": RGBColor(220, 38, 38), "high": RGBColor(234, 88, 12), "medium": RGBColor(245, 158, 11), "low": RGBColor(34, 197, 94)}
                r.font.color.rgb = sev_c.get(finding.get("severity", "medium"), RGBColor(100, 100, 100))
                p.add_run(f"  |  Status: {finding.get('status', 'open').upper()}")

                for label, key in [("Condition", "condition"), ("Criteria", "criteria"), ("Cause", "cause"), ("Effect", "effect")]:
                    if finding.get(key):
                        doc.add_heading(label, level=3)
                        doc.add_paragraph(finding[key])

                if finding.get("owner"):
                    p = doc.add_paragraph()
                    r = p.add_run(f"Owner: {finding['owner']}")
                    r.italic = True

                for mr in finding.get("management_responses", []):
                    doc.add_heading("Management Response", level=3)
                    if mr.get("response_type"):
                        p = doc.add_paragraph()
                        r = p.add_run(f"Response: {mr['response_type'].replace('_', ' ').title()}")
                        r.italic = True
                    if mr.get("response_text"):
                        doc.add_paragraph(mr["response_text"])
                    if mr.get("action_plan"):
                        doc.add_paragraph(f"Action Plan: {mr['action_plan']}")

        section_num += 1
        h = doc.add_heading(f'{section_num}. Recommendations & Action Plans', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(30, 58, 138)

        all_recs = []
        for f in findings:
            for rec in f.get("recommendations", []):
                rec["finding_title"] = f["title"]
                all_recs.append(rec)

        if not all_recs:
            doc.add_paragraph("No recommendations recorded.")
        else:
            for rec in all_recs:
                doc.add_heading(rec["title"], level=2)
                p = doc.add_paragraph()
                r = p.add_run(f"Priority: {rec.get('priority', 'medium').upper()}")
                r.bold = True
                p.add_run(f"  |  Status: {rec.get('status', 'open').upper()}")
                p = doc.add_paragraph()
                r = p.add_run(f"Related Finding: {rec['finding_title']}")
                r.italic = True
                if rec.get("description"):
                    doc.add_paragraph(rec["description"])
                if rec.get("owner"):
                    doc.add_paragraph(f"Owner: {rec['owner']}")
                if rec.get("due_date"):
                    doc.add_paragraph(f"Due Date: {rec['due_date'][:10]}")
                for ap in rec.get("action_plans", []):
                    status_sym = "✓" if ap.get("status") == "completed" else "○"
                    doc.add_paragraph(f"{status_sym} {ap['milestone']} ({ap.get('status', 'pending')})", style='List Bullet')

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = data["title"].replace(" ", "_").replace("/", "_") + ".docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        logger.error(f"DOCX export error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


@router.get("/board-packs")
def list_board_packs(
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"board_packs": [], "total": 0}
    
    packs = db.query(AuditBoardPack).filter(
        AuditBoardPack.tenant_id.in_(user_tenants)
    ).order_by(AuditBoardPack.created_at.desc()).all()
    
    result = []
    for p in packs:
        result.append({
            "id": p.id,
            "title": p.title,
            "period": p.period,
            "executive_summary": p.executive_summary,
            "engagement_ids": p.engagement_ids,
            "key_findings": p.key_findings,
            "kpi_data": p.kpi_data,
            "opinion_summary": p.opinion_summary,
            "status": p.status,
            "ai_generated": p.ai_generated,
            "presented_date": p.presented_date.isoformat() if p.presented_date else None,
            "created_at": p.created_at.isoformat() if p.created_at else None,
        })
    
    return {"board_packs": result, "total": len(result)}


@router.post("/board-packs", status_code=status.HTTP_201_CREATED)
def create_board_pack(
    data: BoardPackCreate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    tenant_id = get_user_primary_tenant(current_user, db)
    if not tenant_id:
        raise HTTPException(status_code=403, detail="No tenant access")
    
    engagement_ids = data.engagement_ids or []
    
    key_findings = []
    kpi_data = {}
    
    if engagement_ids:
        findings = db.query(AuditFinding).filter(
            AuditFinding.engagement_id.in_(engagement_ids),
            AuditFinding.severity.in_(["critical", "high"])
        ).all()
        
        for f in findings:
            key_findings.append({
                "id": f.id,
                "title": f.title,
                "severity": f.severity,
                "status": f.status,
                "engagement_id": f.engagement_id,
            })
        
        total_findings = db.query(func.count(AuditFinding.id)).filter(
            AuditFinding.engagement_id.in_(engagement_ids)
        ).scalar() or 0
        
        closed_findings = db.query(func.count(AuditFinding.id)).filter(
            AuditFinding.engagement_id.in_(engagement_ids),
            AuditFinding.status == "closed"
        ).scalar() or 0
        
        kpi_data = {
            "total_findings": total_findings,
            "closed_findings": closed_findings,
            "closure_rate": round((closed_findings / total_findings * 100) if total_findings > 0 else 0, 1),
            "critical_findings": sum(1 for f in findings if f.severity == "critical"),
            "high_findings": sum(1 for f in findings if f.severity == "high"),
        }
    
    pack = AuditBoardPack(
        tenant_id=tenant_id,
        title=data.title,
        period=data.period,
        executive_summary=data.executive_summary,
        engagement_ids=engagement_ids,
        key_findings=key_findings,
        kpi_data=kpi_data,
        prepared_by_id=current_user.id,
    )
    db.add(pack)
    db.commit()
    db.refresh(pack)
    
    return {
        "id": pack.id,
        "title": pack.title,
        "key_findings": pack.key_findings,
        "kpi_data": pack.kpi_data,
        "status": pack.status,
    }


@router.get("/kpis")
def get_audit_kpis(
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"kpis": {}}
    
    total_findings = db.query(func.count(AuditFinding.id)).filter(
        AuditFinding.tenant_id.in_(user_tenants)
    ).scalar() or 0
    
    closed_findings = db.query(func.count(AuditFinding.id)).filter(
        AuditFinding.tenant_id.in_(user_tenants),
        AuditFinding.status == "closed"
    ).scalar() or 0
    
    overdue_findings = db.query(func.count(AuditFinding.id)).filter(
        AuditFinding.tenant_id.in_(user_tenants),
        AuditFinding.status.in_(["open", "in_progress"]),
        AuditFinding.due_date < datetime.utcnow()
    ).scalar() or 0
    
    total_engagements = db.query(func.count(AuditEngagement.id)).filter(
        AuditEngagement.tenant_id.in_(user_tenants)
    ).scalar() or 0
    
    completed_engagements = db.query(func.count(AuditEngagement.id)).filter(
        AuditEngagement.tenant_id.in_(user_tenants),
        AuditEngagement.status == "closed"
    ).scalar() or 0
    
    active_engagements = db.query(func.count(AuditEngagement.id)).filter(
        AuditEngagement.tenant_id.in_(user_tenants),
        AuditEngagement.status.in_(["planning", "fieldwork", "reporting", "follow_up"])
    ).scalar() or 0
    
    total_budget = db.query(func.sum(AuditEngagement.budget_hours)).filter(
        AuditEngagement.tenant_id.in_(user_tenants)
    ).scalar() or 0
    
    total_actual = db.query(func.sum(AuditEngagement.actual_hours)).filter(
        AuditEngagement.tenant_id.in_(user_tenants)
    ).scalar() or 0
    
    plan_items = db.query(func.count(AuditPlanItem.id)).join(AuditPlan).filter(
        AuditPlan.tenant_id.in_(user_tenants),
        AuditPlan.fiscal_year == str(datetime.utcnow().year)
    ).scalar() or 0
    
    completed_plan_items = db.query(func.count(AuditPlanItem.id)).join(AuditPlan).filter(
        AuditPlan.tenant_id.in_(user_tenants),
        AuditPlan.fiscal_year == str(datetime.utcnow().year),
        AuditPlanItem.status == "completed"
    ).scalar() or 0
    
    closure_rate = (closed_findings / total_findings * 100) if total_findings > 0 else 0
    plan_completion = (completed_plan_items / plan_items * 100) if plan_items > 0 else 0
    cost_efficiency = ((total_budget - total_actual) / total_budget * 100) if total_budget > 0 else 0
    
    return {
        "kpis": {
            "findings_closure_rate": round(closure_rate, 1),
            "overdue_findings": overdue_findings,
            "total_findings": total_findings,
            "active_engagements": active_engagements,
            "completed_engagements": completed_engagements,
            "total_engagements": total_engagements,
            "plan_completion_pct": round(plan_completion, 1),
            "budget_hours": total_budget,
            "actual_hours": total_actual,
            "cost_efficiency_pct": round(cost_efficiency, 1),
        }
    }


@router.get("/trend-analysis")
def get_trend_analysis(
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"trends": {}}
    
    findings = db.query(AuditFinding).filter(
        AuditFinding.tenant_id.in_(user_tenants)
    ).order_by(AuditFinding.created_at.desc()).all()
    
    by_month = {}
    severity_trend = {}
    root_cause_dist = {}
    
    for f in findings:
        if f.created_at:
            month_key = f.created_at.strftime("%Y-%m")
            by_month[month_key] = by_month.get(month_key, 0) + 1
            
            if month_key not in severity_trend:
                severity_trend[month_key] = {}
            severity_trend[month_key][f.severity] = severity_trend[month_key].get(f.severity, 0) + 1
        
        if f.root_cause_category:
            root_cause_dist[f.root_cause_category] = root_cause_dist.get(f.root_cause_category, 0) + 1
    
    return {
        "trends": {
            "findings_by_month": by_month,
            "severity_trends": severity_trend,
            "root_cause_distribution": root_cause_dist,
            "total_findings_analyzed": len(findings),
        }
    }


@router.get("/pbc")
def list_pbc_items(
    engagement_id: Optional[int] = None,
    status_filter: Optional[str] = Query(None, alias="status"),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"pbc_items": [], "total": 0}
    
    query = db.query(PBCListItem).filter(PBCListItem.tenant_id.in_(user_tenants))
    
    if engagement_id:
        query = query.filter(PBCListItem.engagement_id == engagement_id)
    if status_filter:
        query = query.filter(PBCListItem.status == status_filter)
    
    items = query.order_by(PBCListItem.created_at.desc()).all()
    
    result = []
    for item in items:
        result.append({
            "id": item.id,
            "engagement_id": item.engagement_id,
            "document_name": item.document_name,
            "description": item.description,
            "category": item.category,
            "requested_by": item.requested_by,
            "assigned_to_id": item.assigned_to_id,
            "status": item.status,
            "due_date": item.due_date.isoformat() if item.due_date else None,
            "submitted_date": item.submitted_date.isoformat() if item.submitted_date else None,
            "reviewed_date": item.reviewed_date.isoformat() if item.reviewed_date else None,
            "evidence_id": item.evidence_id,
            "notes": item.notes,
            "created_at": item.created_at.isoformat() if item.created_at else None,
        })
    
    return {"pbc_items": result, "total": len(result)}


@router.post("/pbc", status_code=status.HTTP_201_CREATED)
def create_pbc_item(
    data: PBCItemCreate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    tenant_id = get_user_primary_tenant(current_user, db)
    if not tenant_id:
        raise HTTPException(status_code=403, detail="No tenant access")
    
    item = PBCListItem(
        tenant_id=tenant_id,
        engagement_id=data.engagement_id,
        document_name=data.document_name,
        description=data.description,
        category=data.category,
        requested_by=data.requested_by,
        assigned_to_id=data.assigned_to_id,
        due_date=data.due_date,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"id": item.id, "document_name": item.document_name, "status": item.status}


@router.put("/pbc/{item_id}")
def update_pbc_item(
    item_id: int,
    data: PBCItemUpdate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    item = db.query(PBCListItem).filter(
        PBCListItem.id == item_id,
        PBCListItem.tenant_id.in_(user_tenants)
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="PBC item not found")
    
    for field, value in data.dict(exclude_unset=True).items():
        setattr(item, field, value)
    
    if data.status == "submitted":
        item.submitted_date = datetime.utcnow()
    elif data.status == "reviewed":
        item.reviewed_date = datetime.utcnow()
    
    item.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(item)
    return {"id": item.id, "status": item.status, "message": "PBC item updated"}


@router.get("/risk-prioritization")
def get_risk_prioritization(
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"risk_prioritization": {}}

    entities = db.query(AuditableEntity).filter(
        AuditableEntity.tenant_id.in_(user_tenants),
        AuditableEntity.status == "active"
    ).order_by(AuditableEntity.risk_score.desc().nullslast()).all()

    heat_map = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    for e in entities:
        rating = (e.risk_rating or "low").lower()
        if rating in heat_map:
            heat_map[rating] += 1

    risks = db.query(Risk).filter(
        Risk.tenant_id.in_(user_tenants),
        Risk.status != "closed"
    ).all()

    critical_high_risks = [r for r in risks if (r.residual_score or r.inherent_score or 0) >= 50]
    total_critical_high = len(critical_high_risks)

    all_linked_risk_ids = set()
    for e in entities:
        if e.linked_risk_ids:
            for rid in e.linked_risk_ids:
                all_linked_risk_ids.add(rid)

    current_plan = db.query(AuditPlan).filter(
        AuditPlan.tenant_id.in_(user_tenants),
        AuditPlan.fiscal_year == str(datetime.utcnow().year)
    ).first()

    plan_entity_ids = set()
    if current_plan:
        plan_items = db.query(AuditPlanItem).filter(
            AuditPlanItem.audit_plan_id == current_plan.id
        ).all()
        for pi in plan_items:
            if pi.auditable_entity_id:
                plan_entity_ids.add(pi.auditable_entity_id)

    covered_risk_ids = set()
    for e in entities:
        if e.id in plan_entity_ids and e.linked_risk_ids:
            for rid in e.linked_risk_ids:
                covered_risk_ids.add(rid)

    covered_critical_high = len([r for r in critical_high_risks if r.id in covered_risk_ids])
    risk_coverage_pct = round((covered_critical_high / total_critical_high * 100), 1) if total_critical_high > 0 else 100.0

    uncovered = []
    for r in critical_high_risks:
        if r.id not in all_linked_risk_ids:
            uncovered.append({
                "risk_id": r.id,
                "title": r.title,
                "category": r.category,
                "residual_score": r.residual_score,
                "inherent_score": r.inherent_score,
                "status": r.status,
                "owner_id": r.owner_id,
            })

    twelve_months_ago = datetime.utcnow() - timedelta(days=365)
    audited_entities_recent = set()
    recent_engagements = db.query(AuditEngagement).filter(
        AuditEngagement.tenant_id.in_(user_tenants),
        AuditEngagement.status == "closed",
        AuditEngagement.updated_at >= twelve_months_ago
    ).all()
    for eng in recent_engagements:
        if eng.auditable_entity_id:
            audited_entities_recent.add(eng.auditable_entity_id)

    critical_entities = [e for e in entities if (e.risk_rating or "").lower() == "critical"]
    audited_critical = len([e for e in critical_entities if e.id in audited_entities_recent])
    appetite_alignment = round((audited_critical / len(critical_entities) * 100), 1) if critical_entities else 100.0

    now = datetime.utcnow()
    prioritized = []
    for e in entities[:15]:
        last_eng = db.query(AuditEngagement).filter(
            AuditEngagement.auditable_entity_id == e.id,
            AuditEngagement.status == "closed"
        ).order_by(AuditEngagement.updated_at.desc()).first()

        last_audit_date = last_eng.updated_at.isoformat() if last_eng and last_eng.updated_at else None
        days_since = (now - last_eng.updated_at).days if last_eng and last_eng.updated_at else None

        rating = (e.risk_rating or "low").lower()
        if rating == "critical":
            rec_priority = "immediate"
        elif rating == "high":
            rec_priority = "high"
        elif rating == "medium":
            rec_priority = "standard"
        else:
            rec_priority = "low"

        if days_since and days_since > 365:
            rec_priority = "immediate" if rec_priority in ("high", "immediate") else "high"
        elif days_since is None:
            rec_priority = "immediate" if rating in ("critical", "high") else "high"

        prioritized.append({
            "entity_id": e.id,
            "name": e.name,
            "entity_type": e.entity_type,
            "risk_score": e.risk_score,
            "risk_rating": e.risk_rating,
            "last_audit_date": last_audit_date,
            "days_since_audit": days_since,
            "in_current_plan": e.id in plan_entity_ids,
            "recommended_priority": rec_priority,
        })

    return {
        "risk_prioritization": {
            "risk_heat_map": heat_map,
            "risk_coverage_pct": risk_coverage_pct,
            "covered_critical_high": covered_critical_high,
            "total_critical_high": total_critical_high,
            "uncovered_risks": uncovered,
            "risk_appetite_alignment_pct": appetite_alignment,
            "audited_critical_entities": audited_critical,
            "total_critical_entities": len(critical_entities),
            "prioritized_entities": prioritized,
            "plan_risk_alignment_score": current_plan.risk_alignment_score if current_plan else None,
            "risk_scoring_methodology": {
                "model": "Likelihood x Impact",
                "scale": "1-100 composite score",
                "likelihood_factors": ["Historical frequency", "Industry benchmarks", "Control effectiveness", "Environmental factors"],
                "impact_factors": ["Financial loss potential", "Regulatory penalties", "Reputational damage", "Operational disruption"],
                "rating_thresholds": {"critical": ">=75", "high": "50-74", "medium": "25-49", "low": "<25"},
                "frequency_weighting": "Entities overdue for audit receive a 20% score uplift",
                "data_sources": ["Enterprise Risk Register", "Previous audit findings", "Control test results", "Incident history"],
            },
        }
    }


def _build_risk_based_report_data(db: Session, user_tenants, fiscal_year: Optional[str] = None) -> dict:
    """Assemble the board-ready risk-based report: ranked universe, risk-aligned
    plan, and coverage-vs-risk gaps."""
    year = fiscal_year or str(datetime.utcnow().year)
    now = datetime.utcnow()

    entities = db.query(AuditableEntity).filter(
        AuditableEntity.tenant_id.in_(user_tenants),
        AuditableEntity.status == "active",
    ).order_by(AuditableEntity.risk_score.desc().nullslast()).all()

    current_plan = db.query(AuditPlan).filter(
        AuditPlan.tenant_id.in_(user_tenants),
        AuditPlan.fiscal_year == year,
    ).first()

    plan_items = []
    plan_entity_ids = set()
    if current_plan:
        plan_items = db.query(AuditPlanItem).filter(
            AuditPlanItem.audit_plan_id == current_plan.id
        ).all()
        for pi in plan_items:
            if pi.auditable_entity_id:
                plan_entity_ids.add(pi.auditable_entity_id)

    twelve_months_ago = now - timedelta(days=365)
    audited_recent = set()
    recent_engagements = db.query(AuditEngagement).filter(
        AuditEngagement.tenant_id.in_(user_tenants),
        AuditEngagement.status == "closed",
        AuditEngagement.updated_at >= twelve_months_ago,
    ).all()
    for eng in recent_engagements:
        if eng.auditable_entity_id:
            audited_recent.add(eng.auditable_entity_id)

    def _top_factors(entity, limit=3):
        contribs = entity.factor_contributions or []
        if not isinstance(contribs, list):
            return []
        ranked = sorted(
            [c for c in contribs if isinstance(c, dict)],
            key=lambda c: c.get("contribution", 0), reverse=True,
        )
        return [
            {"label": c.get("label"), "value": c.get("value"), "contribution": c.get("contribution")}
            for c in ranked[:limit] if (c.get("contribution") or 0) > 0
        ]

    heat_map = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    ranked_universe = []
    for idx, e in enumerate(entities, start=1):
        rating = (e.risk_rating or "low").lower()
        if rating in heat_map:
            heat_map[rating] += 1
        ranked_universe.append({
            "rank": idx,
            "entity_id": e.id,
            "name": e.name,
            "entity_type": e.entity_type,
            "industry": getattr(e, "industry", None),
            "risk_score": e.risk_score,
            "risk_rating": e.risk_rating,
            "top_factors": _top_factors(e),
            "in_current_plan": e.id in plan_entity_ids,
            "audited_last_12m": e.id in audited_recent,
            "last_audited_date": e.last_audited_date.isoformat() if e.last_audited_date else None,
            "next_audit_due": e.next_audit_due.isoformat() if e.next_audit_due else None,
        })

    entity_by_id = {e.id: e for e in entities}
    plan_rows = []
    for pi in plan_items:
        e = entity_by_id.get(pi.auditable_entity_id)
        plan_rows.append({
            "plan_item_id": pi.id,
            "title": pi.title,
            "entity_id": pi.auditable_entity_id,
            "entity_name": e.name if e else None,
            "risk_score": e.risk_score if e else None,
            "risk_rating": e.risk_rating if e else None,
            "status": pi.status,
            "quarter": getattr(pi, "quarter", None),
            "planned_hours": getattr(pi, "planned_hours", None),
        })
    plan_rows.sort(key=lambda r: (r["risk_score"] or 0), reverse=True)

    high_risk_entities = [e for e in entities if (e.risk_rating or "").lower() in ("critical", "high")]
    coverage_gaps = []
    for e in high_risk_entities:
        if e.id not in plan_entity_ids and e.id not in audited_recent:
            coverage_gaps.append({
                "entity_id": e.id,
                "name": e.name,
                "entity_type": e.entity_type,
                "risk_score": e.risk_score,
                "risk_rating": e.risk_rating,
                "last_audited_date": e.last_audited_date.isoformat() if e.last_audited_date else None,
                "reason": "High/critical risk with no engagement in the last 12 months and not in the current plan",
            })

    total_high = len(high_risk_entities)
    covered_high = len([e for e in high_risk_entities if e.id in plan_entity_ids or e.id in audited_recent])
    coverage_pct = round((covered_high / total_high * 100), 1) if total_high else 100.0

    summary = {
        "fiscal_year": year,
        "total_entities": len(entities),
        "heat_map": heat_map,
        "plan_exists": current_plan is not None,
        "plan_item_count": len(plan_items),
        "high_risk_total": total_high,
        "high_risk_covered": covered_high,
        "high_risk_coverage_pct": coverage_pct,
        "coverage_gap_count": len(coverage_gaps),
        "plan_risk_alignment_score": current_plan.risk_alignment_score if current_plan else None,
    }

    return {
        "generated_at": now.isoformat(),
        "summary": summary,
        "ranked_universe": ranked_universe,
        "risk_aligned_plan": plan_rows,
        "coverage_gaps": coverage_gaps,
    }


@router.get("/risk-based-report")
def get_risk_based_report(
    fiscal_year: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth),
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        raise HTTPException(status_code=403, detail="No tenant access")
    return _build_risk_based_report_data(db, user_tenants, fiscal_year)


@router.get("/risk-based-report/export/docx")
def export_risk_based_report_docx(
    fiscal_year: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth),
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        raise HTTPException(status_code=403, detail="No tenant access")

    data = _build_risk_based_report_data(db, user_tenants, fiscal_year)
    summary = data["summary"]

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading("Risk-Based Audit Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(
            f"Fiscal Year {summary['fiscal_year']} · Generated "
            f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
        )
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("Executive Summary", level=1)
        hm = summary["heat_map"]
        doc.add_paragraph(
            f"The audit universe comprises {summary['total_entities']} active entities: "
            f"{hm['critical']} critical, {hm['high']} high, {hm['medium']} medium and "
            f"{hm['low']} low risk.", style="List Bullet")
        doc.add_paragraph(
            f"High/critical risk coverage stands at {summary['high_risk_coverage_pct']}% "
            f"({summary['high_risk_covered']} of {summary['high_risk_total']} covered by the "
            f"current plan or a recent engagement).", style="List Bullet")
        doc.add_paragraph(
            f"{summary['coverage_gap_count']} high/critical entities are uncovered and represent "
            f"the most material gaps in assurance.", style="List Bullet")
        if summary["plan_risk_alignment_score"] is not None:
            doc.add_paragraph(
                f"Current plan risk-alignment score: {summary['plan_risk_alignment_score']}.",
                style="List Bullet")

        doc.add_heading("Ranked Risk Universe", level=1)
        ranked = data["ranked_universe"][:25]
        if ranked:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for i, h in enumerate(["#", "Entity", "Type", "Score", "Rating"]):
                hdr[i].text = h
            for row in ranked:
                cells = table.add_row().cells
                cells[0].text = str(row["rank"])
                cells[1].text = str(row["name"] or "")
                cells[2].text = str(row["entity_type"] or "")
                cells[3].text = str(round(row["risk_score"]) if row["risk_score"] is not None else "—")
                cells[4].text = str(row["risk_rating"] or "—")
        else:
            doc.add_paragraph("No active entities in the universe.")

        doc.add_heading("Risk-Aligned Audit Plan", level=1)
        plan = data["risk_aligned_plan"]
        if plan:
            ptable = doc.add_table(rows=1, cols=4)
            ptable.style = "Light Grid Accent 1"
            phdr = ptable.rows[0].cells
            for i, h in enumerate(["Audit", "Entity", "Risk", "Status"]):
                phdr[i].text = h
            for row in plan:
                cells = ptable.add_row().cells
                cells[0].text = str(row["title"] or "")
                cells[1].text = str(row["entity_name"] or "—")
                rating = row["risk_rating"] or "—"
                score = round(row["risk_score"]) if row["risk_score"] is not None else "—"
                cells[2].text = f"{rating} ({score})"
                cells[3].text = str(row["status"] or "")
        else:
            doc.add_paragraph(
                f"No audit plan found for fiscal year {summary['fiscal_year']}.")

        doc.add_heading("Coverage vs. Risk Gaps", level=1)
        gaps = data["coverage_gaps"]
        if gaps:
            doc.add_paragraph(
                "The following high/critical risk entities have no coverage in the current "
                "plan and no closed engagement in the last 12 months:")
            for g in gaps:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(
                    f"{g['name']} — {g['risk_rating']} "
                    f"(score {round(g['risk_score']) if g['risk_score'] is not None else '—'})")
                run.bold = True
        else:
            doc.add_paragraph(
                "No coverage gaps detected — all high/critical entities are covered by the "
                "current plan or a recent engagement.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"risk_based_report_{summary['fiscal_year']}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Risk-based report DOCX export error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


@router.get("/accountability")
def get_accountability_metrics(
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"accountability": {}}

    findings = db.query(AuditFinding).options(
        joinedload(AuditFinding.owner),
        joinedload(AuditFinding.engagement),
        joinedload(AuditFinding.management_responses),
    ).filter(
        AuditFinding.tenant_id.in_(user_tenants)
    ).all()

    now = datetime.utcnow()

    owner_map = {}
    for f in findings:
        uid = f.owner_id or 0
        if uid not in owner_map:
            owner_map[uid] = {
                "user_id": uid,
                "display_name": (f.owner.display_name or f.owner.username) if f.owner else "Unassigned",
                "email": f.owner.email if f.owner else None,
                "total": 0,
                "open": 0,
                "overdue": 0,
                "closed": 0,
                "days_to_close": [],
                "oldest_open_days": 0,
            }
        o = owner_map[uid]
        o["total"] += 1
        if f.status == "closed":
            o["closed"] += 1
            if f.created_at and f.updated_at:
                o["days_to_close"].append((f.updated_at - f.created_at).days)
        elif f.status in ("open", "in_progress"):
            o["open"] += 1
            if f.due_date and f.due_date < now:
                o["overdue"] += 1
            if f.created_at:
                age = (now - f.created_at).days
                if age > o["oldest_open_days"]:
                    o["oldest_open_days"] = age

    owner_stats = []
    for uid, o in owner_map.items():
        avg_days = round(sum(o["days_to_close"]) / len(o["days_to_close"]), 1) if o["days_to_close"] else None
        closure_rate = round(o["closed"] / o["total"] * 100, 1) if o["total"] > 0 else 0
        owner_stats.append({
            "user_id": o["user_id"],
            "display_name": o["display_name"],
            "email": o["email"],
            "total_findings": o["total"],
            "open_findings": o["open"],
            "overdue_findings": o["overdue"],
            "closed_findings": o["closed"],
            "closure_rate": closure_rate,
            "avg_days_to_close": avg_days,
            "oldest_open_finding_days": o["oldest_open_days"],
        })
    owner_stats.sort(key=lambda x: x["overdue_findings"], reverse=True)

    escalation_warning = []
    escalation_escalation = []
    escalation_executive = []
    for f in findings:
        if f.status in ("open", "in_progress") and f.due_date and f.due_date < now:
            days_overdue = (now - f.due_date).days
            item = {
                "finding_id": f.id,
                "finding_number": f.finding_number,
                "title": f.title,
                "severity": f.severity,
                "owner_id": f.owner_id,
                "owner_name": (f.owner.display_name or f.owner.username) if f.owner else "Unassigned",
                "engagement_title": f.engagement.title if f.engagement else None,
                "due_date": f.due_date.isoformat(),
                "days_overdue": days_overdue,
            }
            if days_overdue >= 90:
                escalation_executive.append(item)
            elif days_overdue >= 60:
                escalation_escalation.append(item)
            elif days_overdue >= 30:
                escalation_warning.append(item)

    dept_map = {}
    for f in findings:
        dept = "Unknown"
        if f.engagement:
            dept = f.engagement.title.split(" - ")[0] if " - " in f.engagement.title else f.engagement.title
        if dept not in dept_map:
            dept_map[dept] = {"total": 0, "closed": 0, "overdue": 0, "resolution_days": []}
        d = dept_map[dept]
        d["total"] += 1
        if f.status == "closed":
            d["closed"] += 1
            if f.created_at and f.updated_at:
                d["resolution_days"].append((f.updated_at - f.created_at).days)
        if f.status in ("open", "in_progress") and f.due_date and f.due_date < now:
            d["overdue"] += 1

    dept_perf = []
    for dept, d in dept_map.items():
        closure_rate = round(d["closed"] / d["total"] * 100, 1) if d["total"] > 0 else 0
        avg_res = round(sum(d["resolution_days"]) / len(d["resolution_days"]), 1) if d["resolution_days"] else None
        dept_perf.append({
            "department": dept,
            "total_findings": d["total"],
            "closed_findings": d["closed"],
            "overdue_findings": d["overdue"],
            "closure_rate": closure_rate,
            "avg_resolution_days": avg_res,
        })
    dept_perf.sort(key=lambda x: x["overdue_findings"], reverse=True)

    overdue_trend = {}
    for i in range(5, -1, -1):
        month_date = now - timedelta(days=i * 30)
        month_key = month_date.strftime("%Y-%m")
        count = 0
        for f in findings:
            if f.status in ("open", "in_progress") and f.due_date:
                ref_date = datetime(month_date.year, month_date.month, 28)
                if f.due_date < ref_date and (f.status != "closed" or (f.updated_at and f.updated_at > ref_date)):
                    count += 1
        overdue_trend[month_key] = count

    total = len(findings)
    total_closed = sum(1 for f in findings if f.status == "closed")
    total_overdue = sum(1 for f in findings if f.status in ("open", "in_progress") and f.due_date and f.due_date < now)
    total_with_response = sum(1 for f in findings if f.management_responses)

    closure_score = (total_closed / total * 100) if total > 0 else 100
    timeliness_score = max(0, 100 - (total_overdue / max(total, 1) * 200))
    response_score = (total_with_response / total * 100) if total > 0 else 100
    accountability_score = round(closure_score * 0.4 + timeliness_score * 0.3 + response_score * 0.3, 1)

    return {
        "accountability": {
            "accountability_score": min(accountability_score, 100),
            "score_components": {
                "closure_rate_score": round(closure_score, 1),
                "timeliness_score": round(timeliness_score, 1),
                "response_rate_score": round(response_score, 1),
                "weights": {"closure_rate": 0.4, "timeliness": 0.3, "response_rate": 0.3},
            },
            "owner_stats": owner_stats,
            "escalation_summary": {
                "warning": escalation_warning,
                "escalation": escalation_escalation,
                "executive_alert": escalation_executive,
                "total_escalated": len(escalation_warning) + len(escalation_escalation) + len(escalation_executive),
            },
            "department_performance": dept_perf,
            "overdue_trend": overdue_trend,
            "summary": {
                "total_findings": total,
                "total_closed": total_closed,
                "total_overdue": total_overdue,
                "total_with_management_response": total_with_response,
            },
        }
    }


@router.get("/roi-metrics")
def get_roi_metrics(
    hourly_rate: float = Query(150.0, description="Hourly rate for cost calculations"),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"roi_metrics": {}}

    findings = db.query(AuditFinding).options(
        joinedload(AuditFinding.engagement),
        joinedload(AuditFinding.recommendations),
    ).filter(
        AuditFinding.tenant_id.in_(user_tenants)
    ).all()

    engagements = db.query(AuditEngagement).filter(
        AuditEngagement.tenant_id.in_(user_tenants)
    ).all()

    total_impact = sum(f.estimated_financial_impact or 0 for f in findings)
    total_savings = sum(f.actual_savings_realized or 0 for f in findings if f.status in ("remediated", "closed"))

    total_actual_hours = sum(e.actual_hours or 0 for e in engagements)
    total_audit_cost = total_actual_hours * hourly_rate
    roi_ratio = round(total_savings / total_audit_cost, 2) if total_audit_cost > 0 else 0

    impact_by_category = {}
    for f in findings:
        cat = f.impact_category or "unclassified"
        if cat not in impact_by_category:
            impact_by_category[cat] = {"count": 0, "estimated_impact": 0, "realized_savings": 0}
        impact_by_category[cat]["count"] += 1
        impact_by_category[cat]["estimated_impact"] += f.estimated_financial_impact or 0
        impact_by_category[cat]["realized_savings"] += f.actual_savings_realized or 0

    eng_map = {}
    for f in findings:
        eid = f.engagement_id
        if eid not in eng_map:
            eng_map[eid] = {
                "engagement_id": eid,
                "title": f.engagement.title if f.engagement else "Unknown",
                "findings_count": 0,
                "total_impact": 0,
                "savings_realized": 0,
            }
        eng_map[eid]["findings_count"] += 1
        eng_map[eid]["total_impact"] += f.estimated_financial_impact or 0
        eng_map[eid]["savings_realized"] += f.actual_savings_realized or 0

    value_by_engagement = []
    for eid, em in eng_map.items():
        eng_obj = next((e for e in engagements if e.id == eid), None)
        hours = eng_obj.actual_hours or 0 if eng_obj else 0
        cost = hours * hourly_rate
        eng_roi = round(em["savings_realized"] / cost, 2) if cost > 0 else 0
        value_by_engagement.append({
            **em,
            "hours_invested": hours,
            "cost": round(cost, 2),
            "roi": eng_roi,
        })
    value_by_engagement.sort(key=lambda x: x["total_impact"], reverse=True)

    all_recs = []
    for f in findings:
        all_recs.extend(f.recommendations or [])

    total_recs = len(all_recs)
    recs_with_value = sum(1 for r in all_recs if r.estimated_value)
    recs_implemented = sum(1 for r in all_recs if r.status in ("implemented", "closed", "completed"))
    total_est_value = sum(r.estimated_value or 0 for r in all_recs)
    total_real_value = sum(r.realized_value or 0 for r in all_recs)

    total_findings = len(findings)
    total_engs = len(engagements)
    avg_hours_per_finding = round(total_actual_hours / total_findings, 1) if total_findings > 0 else 0
    avg_cost_per_finding = round(total_audit_cost / total_findings, 2) if total_findings > 0 else 0
    findings_per_engagement = round(total_findings / total_engs, 1) if total_engs > 0 else 0
    cost_per_engagement = round(total_audit_cost / total_engs, 2) if total_engs > 0 else 0

    return {
        "roi_metrics": {
            "total_financial_impact_identified": round(total_impact, 2),
            "total_savings_realized": round(total_savings, 2),
            "total_audit_cost": round(total_audit_cost, 2),
            "roi_ratio": roi_ratio,
            "hourly_rate_used": hourly_rate,
            "impact_by_category": impact_by_category,
            "value_by_engagement": value_by_engagement[:20],
            "recommendation_value_tracking": {
                "total_recommendations": total_recs,
                "with_value_estimate": recs_with_value,
                "implemented": recs_implemented,
                "total_estimated_value": round(total_est_value, 2),
                "total_realized_value": round(total_real_value, 2),
                "realization_rate": round(total_real_value / total_est_value * 100, 1) if total_est_value > 0 else 0,
            },
            "efficiency_metrics": {
                "total_findings": total_findings,
                "total_engagements": total_engs,
                "total_audit_hours": total_actual_hours,
                "total_audit_cost": round(total_audit_cost, 2),
                "avg_hours_per_finding": avg_hours_per_finding,
                "avg_cost_per_finding": avg_cost_per_finding,
                "findings_per_engagement": findings_per_engagement,
                "cost_per_engagement": cost_per_engagement,
            },
        }
    }


@router.get("/regulatory-impact-tracker")
def get_regulatory_impact_tracker(
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        raise HTTPException(status_code=403, detail="No tenant access")

    now = datetime.utcnow()
    ninety_days_ago = now - timedelta(days=90)

    active_changes = db.query(RegulatoryChange).filter(
        RegulatoryChange.tenant_id.in_(user_tenants),
        RegulatoryChange.status.in_(["identified", "under_assessment", "implementation"])
    ).order_by(RegulatoryChange.priority.asc(), RegulatoryChange.effective_date.asc()).all()

    completed_changes = db.query(RegulatoryChange).filter(
        RegulatoryChange.tenant_id.in_(user_tenants),
        RegulatoryChange.status.in_(["completed", "closed"]),
        RegulatoryChange.updated_at >= ninety_days_ago
    ).order_by(RegulatoryChange.updated_at.desc()).limit(10).all()

    active_plan_items = db.query(AuditPlanItem).join(AuditPlan).filter(
        AuditPlan.tenant_id.in_(user_tenants),
        AuditPlan.status.in_(["draft", "approved", "active"])
    ).all()

    active_engagements = db.query(AuditEngagement).filter(
        AuditEngagement.tenant_id.in_(user_tenants),
        AuditEngagement.status.in_(["planning", "fieldwork", "reporting"])
    ).all()

    plan_item_names = set()
    for pi in active_plan_items:
        plan_item_names.add((pi.name or "").lower())
    engagement_titles = set()
    for eng in active_engagements:
        engagement_titles.add((eng.title or "").lower())

    priority_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}

    pending_changes = []
    changes_requiring_audit = []
    covered_count = 0

    for rc in active_changes:
        gap_count = 0
        assessments = db.query(RegulatoryImpactAssessment).filter(
            RegulatoryImpactAssessment.regulatory_change_id == rc.id
        ).all()
        for a in assessments:
            if a.gap_identified:
                gap_count += 1

        days_until_effective = None
        if rc.effective_date:
            delta = rc.effective_date - now
            days_until_effective = delta.days

        rc_title_lower = (rc.title or "").lower()
        rc_source_lower = (rc.source or "").lower()
        has_coverage = False
        for pn in plan_item_names:
            if pn and (rc_title_lower in pn or pn in rc_title_lower or rc_source_lower in pn):
                has_coverage = True
                break
        if not has_coverage:
            for et in engagement_titles:
                if et and (rc_title_lower in et or et in rc_title_lower or rc_source_lower in et):
                    has_coverage = True
                    break

        audit_impact = None
        if hasattr(rc, 'audit_impact_analysis') and rc.audit_impact_analysis:
            audit_impact = rc.audit_impact_analysis

        change_data = {
            "id": rc.id,
            "title": rc.title,
            "description": (rc.description or "")[:200],
            "source": rc.source,
            "priority": rc.priority,
            "status": rc.status,
            "effective_date": rc.effective_date.isoformat() if rc.effective_date else None,
            "days_until_effective": days_until_effective,
            "gap_count": gap_count,
            "has_audit_coverage": has_coverage,
            "audit_impact_analysis": audit_impact,
        }
        pending_changes.append(change_data)

        if has_coverage:
            covered_count += 1
        else:
            if rc.priority in ("critical", "high") or gap_count > 0:
                action = "new_engagement"
                rationale = f"No audit coverage found for this {rc.priority}-priority regulatory change."
                if days_until_effective is not None and days_until_effective < 90:
                    rationale += f" Only {days_until_effective} days until effective date — urgent action needed."
                    action = "new_engagement"
                elif days_until_effective is not None and days_until_effective < 180:
                    rationale += f" {days_until_effective} days until effective date — plan scope expansion."
                    action = "scope_expansion"

                if gap_count > 0:
                    rationale += f" {gap_count} gap(s) identified in impact assessment."

                changes_requiring_audit.append({
                    **change_data,
                    "recommended_action": action,
                    "rationale": rationale,
                })

    pending_changes.sort(key=lambda x: (priority_order.get(x["priority"], 3), x["days_until_effective"] or 9999))

    total_active = len(active_changes)
    coverage_pct = round(covered_count / total_active * 100, 1) if total_active > 0 else 100.0

    recently_completed = []
    for rc in completed_changes:
        recently_completed.append({
            "id": rc.id,
            "title": rc.title,
            "source": rc.source,
            "priority": rc.priority,
            "status": rc.status,
            "completed_at": rc.updated_at.isoformat() if rc.updated_at else None,
        })

    pending_new = db.query(func.count(RegulatoryFeedItem.id)).filter(
        RegulatoryFeedItem.tenant_id.in_(user_tenants),
        RegulatoryFeedItem.status == "new"
    ).scalar() or 0
    pending_analyzed = db.query(func.count(RegulatoryFeedItem.id)).filter(
        RegulatoryFeedItem.tenant_id.in_(user_tenants),
        RegulatoryFeedItem.status == "analyzed",
        RegulatoryFeedItem.regulatory_change_id == None
    ).scalar() or 0

    return {
        "regulatory_impact_tracker": {
            "impact_summary": {
                "total_active_changes": total_active,
                "changes_with_audit_coverage": covered_count,
                "changes_without_coverage": total_active - covered_count,
                "coverage_percentage": coverage_pct,
            },
            "pending_changes": pending_changes,
            "changes_requiring_audit": changes_requiring_audit,
            "audit_actions_needed": changes_requiring_audit,
            "recently_completed": recently_completed,
            "pending_feed_items": {
                "new_count": pending_new,
                "analyzed_count": pending_analyzed,
                "total_unprocessed": pending_new + pending_analyzed,
            },
        }
    }


@router.get("/storyboard-pdf")
def generate_audit_storyboard_pdf(
    current_user: GRCUser = Depends(require_auth)
):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=30*mm, bottomMargin=25*mm, leftMargin=25*mm, rightMargin=25*mm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=28, spaceAfter=6, textColor=colors.HexColor('#1e293b'), fontName='Helvetica-Bold')
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=14, spaceAfter=20, textColor=colors.HexColor('#64748b'), fontName='Helvetica')
    heading_style = ParagraphStyle('ModuleHeading', parent=styles['Heading1'], fontSize=18, spaceBefore=16, spaceAfter=8, textColor=colors.HexColor('#1e40af'), fontName='Helvetica-Bold')
    subheading_style = ParagraphStyle('SubHeading', parent=styles['Heading2'], fontSize=13, spaceBefore=10, spaceAfter=4, textColor=colors.HexColor('#334155'), fontName='Helvetica-Bold')
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, spaceAfter=6, textColor=colors.HexColor('#334155'), leading=14)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'], fontSize=10, spaceAfter=3, textColor=colors.HexColor('#475569'), leftIndent=18, bulletIndent=6, leading=14)
    ai_badge_style = ParagraphStyle('AIBadge', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#7c3aed'), fontName='Helvetica-BoldOblique')
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.HexColor('#94a3b8'))

    story = []

    story.append(Spacer(1, 40))
    story.append(Paragraph("AuditVerse.AI", title_style))
    story.append(Paragraph("Audit Management Module — Product Storyboard", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#3b82f6'), spaceAfter=12))
    story.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}", footer_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "A comprehensive, AI-powered audit lifecycle management system following IIA 2024 Global Internal Audit Standards. "
        "This module provides end-to-end coverage from audit universe management through engagement execution, findings tracking, "
        "continuous monitoring, and quality assurance — all integrated with the broader GRC platform.",
        body_style
    ))
    story.append(Spacer(1, 8))

    overview_data = [
        ['Metric', 'Value'],
        ['Total Pages/Modules', '11'],
        ['AI-Powered Features', '12+'],
        ['Database Models', '24'],
        ['Backend Routers', '12'],
        ['Standards Alignment', 'IIA 2024 GIAS'],
    ]
    overview_table = Table(overview_data, colWidths=[200, 200])
    overview_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(overview_table)
    story.append(PageBreak())

    modules = [
        {
            "title": "1. Audit Dashboard",
            "purpose": "Central command center providing real-time visibility into the entire audit function's performance, risk posture, and regulatory compliance status.",
            "features": [
                "KPI cards: open engagements, overdue findings, average cycle time, risk coverage score",
                "Regulatory Change Impact tracker with coverage percentage gauge",
                "Pending regulatory feed items banner with direct links to process them",
                "Risk prioritization matrix and trend charts",
                "ROI metrics: cost savings, efficiency gains, compliance penalties avoided",
                "Accountability dashboard with team performance metrics",
                "Alert banners for critical/high uncovered regulatory changes",
            ],
            "ai": "AI-powered regulatory impact assessment, automated coverage gap detection",
            "workflow": "Auditors land here to see what needs attention → drill into specific modules from dashboard cards",
        },
        {
            "title": "2. Audit Universe",
            "purpose": "Comprehensive inventory of all auditable entities across the organization, with risk-based prioritization to drive audit planning.",
            "features": [
                "Entity management with risk scores (inherent, residual, control effectiveness)",
                "Coverage gap analysis — identifies entities that haven't been audited recently",
                "Risk Register integration: sync entities from risk categories and individual high-severity risks",
                "Refresh risk scores to keep audit universe current with latest risk assessments",
                "Entity categorization by type: process, department, system, third-party, project",
                "Linked risk IDs for bi-directional risk tracing",
            ],
            "ai": "AI-powered risk scoring suggestions, automated coverage gap identification",
            "workflow": "Maintain universe → sync from risks → identify gaps → feed into audit planning",
        },
        {
            "title": "3. Audit Plans",
            "purpose": "Annual and multi-year audit planning with risk-based resource allocation, regulatory change awareness, and approval workflows.",
            "features": [
                "Create plans manually, from universe coverage gaps, or AI-generated",
                "Plan items with scope, budget days, assigned auditors, risk alignment scores",
                "Approval workflow: draft → submitted → approved",
                "Generate from Universe pipeline: auto-creates plan items from highest-risk entities",
                "Regulatory Change Impact section showing changes needing audit coverage",
                "Bulk create engagements from approved plan items",
                "Calendar view for scheduling",
                "Notification badges for regulatory changes needing attention",
            ],
            "ai": "AI Generate Plan: creates a complete risk-based audit plan with scope, timing, and resource allocation",
            "workflow": "Define plan → add items from universe/AI → get approval → create engagements",
        },
        {
            "title": "4. Engagements",
            "purpose": "Execution management for individual audit engagements, tracking the full lifecycle from planning through fieldwork to reporting and follow-up.",
            "features": [
                "Lifecycle pipeline: planning → fieldwork → reporting → follow_up → closed",
                "Team management: lead auditor, team members, specialist assignments",
                "Time tracking and budget monitoring with variance analysis",
                "Create from plan items (single or bulk)",
                "Workpaper management with three-level sign-off (preparer → reviewer → lead)",
                "Evidence cross-referencing with the Evidence Management module",
                "Status-specific action buttons and progress tracking",
            ],
            "ai": "AI Fieldwork Guidance for planning/fieldwork phases; AI Risk Assessment Suggestions for reporting/follow-up phases; AI-generated engagement details from plan context",
            "workflow": "Create from plan → assign team → execute fieldwork → document workpapers → draft findings → report → follow up → close",
        },
        {
            "title": "5. Findings",
            "purpose": "Structured management of audit findings using the Condition/Criteria/Cause/Effect (CCCE) format with management responses, action plans, and follow-up tracking.",
            "features": [
                "CCCE format: Condition, Criteria, Cause, Effect for each finding",
                "Severity levels: critical, high, medium, low, advisory",
                "Management responses and agreed action plans with due dates",
                "Framework mappings: SOX, COSO, IIA, IFRS, ISO 27001",
                "Follow-up tracking with evidence of remediation",
                "Financial impact tracking: estimated impact, actual savings, impact category",
                "Critical/high findings auto-create Risk Register entries",
                "Recommendations linked to findings with implementation status",
            ],
            "ai": "AI Draft Finding: auto-populates CCCE fields from a brief description; AI Similarity Detection: checks for similar past findings with percentage match scores",
            "workflow": "Identify issue → draft finding (AI optional) → assign severity → get management response → track action plan → verify remediation → close",
        },
        {
            "title": "6. Continuous Control Monitoring (CCM)",
            "purpose": "Automated, rule-based monitoring of control effectiveness with anomaly detection and exception management workflows.",
            "features": [
                "Monitoring rules: threshold, pattern, statistical, ML-based detection types",
                "Real-time anomaly detection with severity classification",
                "Exception workflow: detected → investigating → resolved / escalated",
                "Anomaly-to-finding escalation pipeline",
                "Control coverage statistics and effectiveness metrics",
                "Rule configuration with custom parameters and thresholds",
                "Dashboard showing active monitors, recent anomalies, and trends",
            ],
            "ai": "AI CCM Insights: analyzes patterns in detected anomalies and suggests rule optimizations",
            "workflow": "Configure rules → monitor continuously → detect anomalies → investigate → escalate to findings if needed",
        },
        {
            "title": "7. Reporting",
            "purpose": "Comprehensive audit reporting with engagement reports, executive summaries, board packs, and AI-generated narratives for stakeholder communication.",
            "features": [
                "KPI trends: finding closure rate, engagement completion, overdue items",
                "Engagement reports with opinions (satisfactory, needs improvement, unsatisfactory)",
                "Full-page professional report view with PDF and DOCX export",
                "Board packs with executive summaries and key metrics aggregation",
                "Report lifecycle: draft → review → issued",
                "ROI metrics and financial impact summaries",
                "Regulatory impact tracker integration",
            ],
            "ai": "AI Board Pack Narrative: generates professional executive summaries from engagement data; AI-generated report content from findings and recommendations",
            "workflow": "Create report from engagement → add executive summary → set opinion → generate board pack → export PDF/DOCX → distribute to stakeholders",
        },
        {
            "title": "8. Quality Assurance & Improvement (QAIP)",
            "purpose": "Internal audit quality management following IIA standards, with conformance scoring, maturity assessments, and continuous improvement tracking.",
            "features": [
                "Internal and external quality reviews",
                "IIA conformance scoring against Global Internal Audit Standards",
                "5-level maturity model: Initial → Developing → Defined → Managed → Optimizing",
                "Audit templates for standardized engagement execution",
                "Improvement action tracking with due dates and ownership",
                "Quality metrics dashboard",
            ],
            "ai": None,
            "workflow": "Schedule reviews → assess conformance → score maturity → identify improvements → track remediation → reassess",
        },
        {
            "title": "9. Test Script Library",
            "purpose": "Centralized repository of reusable test scripts with structured procedures, enabling consistent and efficient control testing across engagements.",
            "features": [
                "CRUD management of reusable test scripts",
                "Control area categorization for easy discovery",
                "Structured test steps with expected results",
                "Clone test scripts into engagements as workpapers",
                "Usage tracking (last used date, engagement count)",
                "Entity type tagging for contextual filtering",
            ],
            "ai": "AI Generate from Engagement: auto-creates test scripts based on engagement context, scope, findings, and control objectives",
            "workflow": "Create/browse scripts → select for engagement → clone to workpapers → execute during fieldwork → update library with lessons learned",
        },
        {
            "title": "10. Auditor Skill Matrix",
            "purpose": "Resource competency management tracking auditor skills, certifications, and proficiency levels to enable optimal team assignment and capacity planning.",
            "features": [
                "Skill proficiency tracking (1-5 scale) across audit categories",
                "Certification management with expiry date alerts",
                "Find-auditor matching: search by required skills and framework expertise",
                "Team competency statistics and gap analysis",
                "8 audit category taxonomy: Financial, IT/Cybersecurity, Operational, Banking, Compliance, Forensic, ESG, Internal Audit",
                "L1/L2/L3 skill hierarchy for granular competency mapping",
            ],
            "ai": "AI Suggest Skills: recommends exhaustive skill taxonomy with hierarchical categorization across all audit domains",
            "workflow": "Define skill requirements → assess auditor proficiencies → match skills to engagements → identify training needs → track certification renewals",
        },
        {
            "title": "11. Capacity Planning",
            "purpose": "Auditor workload management with allocation calendars, utilization analytics, and conflict detection to prevent over-commitment and ensure adequate coverage.",
            "features": [
                "Auditor allocation calendar view",
                "Utilization statistics: allocated vs available hours",
                "Conflict detection for over-allocated auditors",
                "Engagement-level hour allocation and tracking",
                "Team-wide capacity overview",
                "Availability checking before assignment",
            ],
            "ai": None,
            "workflow": "View team capacity → check availability → allocate to engagements → monitor utilization → resolve conflicts → rebalance workload",
        },
    ]

    for i, mod in enumerate(modules):
        story.append(Paragraph(mod["title"], heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#93c5fd'), spaceAfter=8))

        story.append(Paragraph("<b>Purpose</b>", subheading_style))
        story.append(Paragraph(mod["purpose"], body_style))

        story.append(Paragraph("<b>Key Features</b>", subheading_style))
        for feat in mod["features"]:
            story.append(Paragraph(f"\u2022  {feat}", bullet_style))

        if mod.get("ai"):
            story.append(Spacer(1, 4))
            story.append(Paragraph("\u2728 AI-Powered Capabilities", ai_badge_style))
            story.append(Paragraph(mod["ai"], body_style))

        story.append(Paragraph("<b>User Workflow</b>", subheading_style))
        story.append(Paragraph(mod["workflow"], body_style))

        if i < len(modules) - 1:
            story.append(PageBreak())

    story.append(PageBreak())
    story.append(Paragraph("Integration Points", heading_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#93c5fd'), spaceAfter=8))

    integrations = [
        ["Integration", "Description"],
        ["Risk Register", "Critical/high findings auto-create risk entries; Audit Universe syncs from risk categories"],
        ["Evidence Management", "Workpapers cross-reference Evidence items; engagement evidence linked bi-directionally"],
        ["Regulatory Changes", "Feed items analyzed and converted to track audit coverage; impact tracker on dashboard and plans"],
        ["Compliance Frameworks", "Findings mapped to SOX/COSO/IIA/IFRS/ISO27001; framework-aware test scripts"],
        ["Governance Policies", "Policy attestation tracking; regulatory change impact on policies surfaced"],
    ]
    int_table = Table(integrations, colWidths=[140, 310])
    int_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.append(int_table)
    story.append(Spacer(1, 20))

    story.append(Paragraph("AI Capabilities Summary", heading_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#93c5fd'), spaceAfter=8))

    ai_data = [
        ["AI Feature", "Module", "Trigger"],
        ["Generate Audit Plan", "Audit Plans", "Button click"],
        ["Generate Procedures", "Engagements", "Button click"],
        ["Draft Finding (CCCE)", "Findings", "Button click"],
        ["Similarity Detection", "Findings", "On create/edit"],
        ["CCM Insights", "CCM", "Button click"],
        ["Board Pack Narrative", "Reporting", "Button click"],
        ["Engagement Details", "Engagements", "Button click"],
        ["Risk Assessment Suggestions", "Engagements", "Status-based"],
        ["Fieldwork Guidance", "Engagements", "Status-based"],
        ["Regulatory Impact Assessment", "Plans/Dashboard", "Auto on conversion"],
        ["Generate Test Script", "Test Scripts", "Button click"],
        ["Suggest Auditor Skills", "Skill Matrix", "Button click"],
    ]
    ai_table = Table(ai_data, colWidths=[170, 120, 160])
    ai_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f5f3ff')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#c4b5fd')),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(ai_table)

    story.append(Spacer(1, 30))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0'), spaceAfter=8))
    story.append(Paragraph("AuditVerse.AI \u2014 Enterprise Audit Platform", footer_style))
    story.append(Paragraph(f"Audit Management Module Storyboard \u2022 {datetime.utcnow().strftime('%Y')}", footer_style))

    doc.build(story)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=AuditVerse_AI_Audit_Management_Storyboard.pdf"}
    )
